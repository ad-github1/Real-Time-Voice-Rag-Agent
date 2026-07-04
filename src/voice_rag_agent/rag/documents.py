from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable


SOURCE_SUFFIXES = {".md", ".markdown", ".txt", ".pdf", ".docx", ".html", ".htm"}


@dataclass(frozen=True)
class DocumentChunk:
    id: str
    source_path: str
    text: str
    metadata: dict[str, str | int]


def iter_source_files(root: Path) -> Iterable[Path]:
    if not root.exists():
        return []
    return sorted(
        path
        for path in root.rglob("*")
        if path.is_file() and path.suffix.lower() in SOURCE_SUFFIXES
    )


def extract_title(text: str, fallback: str) -> str:
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            return stripped[2:].strip()
        if stripped:
            return stripped[:80]
    return fallback


def split_text(text: str, max_chars: int, overlap: int) -> list[str]:
    cleaned = re.sub(r"\n{3,}", "\n\n", text.strip())
    cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)

    if not cleaned:
        return []

    paragraphs = [paragraph.strip() for paragraph in cleaned.split("\n\n") if paragraph.strip()]
    chunks: list[str] = []
    current = ""

    for paragraph in paragraphs:
        if len(paragraph) > max_chars:
            if current:
                chunks.append(current.strip())
                current = ""
            chunks.extend(_split_long_paragraph(paragraph, max_chars, overlap))
            continue

        candidate = f"{current}\n\n{paragraph}".strip() if current else paragraph

        if len(candidate) <= max_chars:
            current = candidate
        else:
            chunks.append(current.strip())
            prefix = _tail(current, overlap)
            current = f"{prefix}\n\n{paragraph}".strip() if prefix else paragraph

    if current:
        chunks.append(current.strip())

    return chunks


def _split_long_paragraph(paragraph: str, max_chars: int, overlap: int) -> list[str]:
    words = paragraph.split()
    chunks: list[str] = []
    current: list[str] = []

    for word in words:
        candidate = " ".join([*current, word])

        if len(candidate) <= max_chars:
            current.append(word)
            continue

        if current:
            chunk = " ".join(current)
            chunks.append(chunk)
            current = _tail(chunk, overlap).split()

        current.append(word)

    if current:
        chunks.append(" ".join(current))

    return chunks


def _tail(text: str, max_chars: int) -> str:
    if max_chars <= 0:
        return ""

    if len(text) <= max_chars:
        return text

    tail = text[-max_chars:].lstrip()
    first_space = tail.find(" ")

    return tail[first_space + 1 :] if first_space > 0 else tail


def read_source_text(source: Path) -> str:
    suffix = source.suffix.lower()

    if suffix in {".md", ".markdown", ".txt"}:
        return source.read_text(encoding="utf-8", errors="ignore")

    if suffix == ".pdf":
        return _read_pdf(source)

    if suffix == ".docx":
        return _read_docx(source)

    if suffix in {".html", ".htm"}:
        html = source.read_text(encoding="utf-8", errors="ignore")
        return _html_to_text(html)

    return ""


def _read_pdf(source: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "PDF ingestion needs pypdf. Install it with: python3 -m pip install pypdf"
        ) from exc

    reader = PdfReader(str(source))
    pages: list[str] = []

    for index, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text = text.strip()

        if text:
            pages.append(f"\n\n[Page {index + 1}]\n{text}")

    return "\n".join(pages)


def _read_docx(source: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "DOCX ingestion needs python-docx. Install it with: python3 -m pip install python-docx"
        ) from exc

    document = Document(str(source))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def _html_to_text(html: str) -> str:
    try:
        from bs4 import BeautifulSoup
    except ImportError as exc:
        raise RuntimeError(
            "HTML ingestion needs beautifulsoup4. Install it with: python3 -m pip install beautifulsoup4"
        ) from exc

    soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["script", "style", "noscript", "svg"]):
        tag.decompose()

    title = ""
    if soup.title and soup.title.string:
        title = f"# {soup.title.string.strip()}\n\n"

    text = soup.get_text(separator="\n")
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    return title + "\n".join(lines)


def iter_web_sources(root: Path) -> Iterable[tuple[str, str]]:
    """Read URLs from data_dir/urls.txt and fetch them as web documents.

    Each non-empty, non-comment line in urls.txt should be one URL.
    """

    urls_file = root / "urls.txt"

    if not urls_file.exists():
        return []

    urls = [
        line.strip()
        for line in urls_file.read_text(encoding="utf-8", errors="ignore").splitlines()
        if line.strip() and not line.strip().startswith("#")
    ]

    return [_fetch_url(url) for url in urls]


def _fetch_url(url: str) -> tuple[str, str]:
    try:
        import httpx
    except ImportError as exc:
        raise RuntimeError(
            "Web ingestion needs httpx. Install it with: python3 -m pip install httpx"
        ) from exc

    response = httpx.get(
        url,
        timeout=15.0,
        follow_redirects=True,
        headers={"User-Agent": "voice-rag-agent/0.1"},
    )
    response.raise_for_status()

    content_type = response.headers.get("content-type", "").lower()

    if "html" in content_type:
        return url, _html_to_text(response.text)

    return url, response.text


def _add_chunks(
    chunks: list[DocumentChunk],
    *,
    source_path: str,
    text: str,
    fallback_title: str,
    chunk_size: int,
    chunk_overlap: int,
) -> None:
    title = extract_title(text, fallback_title)

    for index, chunk_text in enumerate(split_text(text, chunk_size, chunk_overlap)):
        digest = hashlib.sha1(f"{source_path}:{index}:{chunk_text}".encode("utf-8")).hexdigest()[
            :12
        ]

        chunks.append(
            DocumentChunk(
                id=f"chunk_{digest}",
                source_path=source_path,
                text=chunk_text,
                metadata={
                    "title": title,
                    "path": source_path,
                    "chunk_index": index,
                },
            )
        )


def load_documents(
    root: Path, chunk_size: int = 900, chunk_overlap: int = 120
) -> list[DocumentChunk]:
    chunks: list[DocumentChunk] = []

    for source in iter_source_files(root):
        text = read_source_text(source)

        if not text.strip():
            continue

        _add_chunks(
            chunks,
            source_path=str(source),
            text=text,
            fallback_title=source.stem.replace("_", " ").title(),
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )

    for url, text in iter_web_sources(root):
        if not text.strip():
            continue

        _add_chunks(
            chunks,
            source_path=url,
            text=text,
            fallback_title=url,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )

    return chunks
